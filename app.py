import streamlit as st
import pandas as pd
from docx import Document
import subprocess
import platform
from PyPDF2 import PdfMerger
import io
import re
import tempfile
import os
import base64
import zipfile
from openpyxl import load_workbook

def load_excel_file(uploaded_file):
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
            tmp.write(uploaded_file.getvalue())
            tmp_path = tmp.name

        wb = load_workbook(tmp_path, data_only=True)
        ws = wb.active
        headers = [cell.value for cell in ws[1]]

        data = []
        for row in ws.iter_rows(min_row=2):
            row_data = {}
            for header, cell in zip(headers, row):
                if cell.value is None:
                    row_data[header] = ""
                else:
                    if cell.is_date:
                        row_data[header] = cell.value.strftime("%d/%m/%Y")
                    else:
                        row_data[header] = str(cell.value)
            data.append(row_data)

        os.unlink(tmp_path)
        return pd.DataFrame(data)

    except Exception as e:
        st.error(f"Lỗi khi đọc file Excel: {str(e)}")
        return None

def replace_placeholders_in_paragraph(paragraph, data_dict):
    for key, value in data_dict.items():
        placeholder = f"{{{{{key}}}}}"
        full_text = ''.join(run.text for run in paragraph.runs)
        
        if placeholder not in full_text:
            continue
        
        match = re.search(re.escape(placeholder), full_text)
        if not match:
            continue
            
        placeholder_start = match.start()
        placeholder_end = match.end()
        
        current_pos = 0
        placeholder_format = None
        
        for run in paragraph.runs:
            run_length = len(run.text)
            run_end = current_pos + run_length
            
            if current_pos <= placeholder_start < run_end:
                placeholder_format = {
                    'font_name': run.font.name,
                    'font_size': run.font.size,
                    'bold': run.font.bold,
                    'italic': run.font.italic,
                    'underline': run.font.underline,
                    'color': run.font.color.rgb if run.font.color.rgb else None,
                    'highlight': run.font.highlight_color
                }
                break
            current_pos = run_end
        
        if not placeholder_format:
            continue
        
        new_runs = []
        current_pos = 0
        
        for run in paragraph.runs:
            run_text = run.text
            run_start = current_pos
            run_end = current_pos + len(run_text)
            
            if run_end <= placeholder_start:
                new_runs.append({
                    'text': run_text,
                    'format': {
                        'font_name': run.font.name,
                        'font_size': run.font.size,
                        'bold': run.font.bold,
                        'italic': run.font.italic,
                        'underline': run.font.underline,
                        'color': run.font.color.rgb if run.font.color.rgb else None,
                        'highlight': run.font.highlight_color
                    }
                })
            elif run_start >= placeholder_end:
                new_runs.append({
                    'text': run_text,
                    'format': {
                        'font_name': run.font.name,
                        'font_size': run.font.size,
                        'bold': run.font.bold,
                        'italic': run.font.italic,
                        'underline': run.font.underline,
                        'color': run.font.color.rgb if run.font.color.rgb else None,
                        'highlight': run.font.highlight_color
                    }
                })
            else:
                if run_start < placeholder_start:
                    before_text = run_text[:placeholder_start - run_start]
                    new_runs.append({
                        'text': before_text,
                        'format': {
                            'font_name': run.font.name,
                            'font_size': run.font.size,
                            'bold': run.font.bold,
                            'italic': run.font.italic,
                            'underline': run.font.underline,
                            'color': run.font.color.rgb if run.font.color.rgb else None,
                            'highlight': run.font.highlight_color
                        }
                    })
                
                if run_start <= placeholder_start < run_end:
                    new_runs.append({
                        'text': str(value),
                        'format': placeholder_format
                    })
                
                if run_end > placeholder_end:
                    after_start_in_run = max(0, placeholder_end - run_start)
                    after_text = run_text[after_start_in_run:]
                    if after_text:
                        new_runs.append({
                            'text': after_text,
                            'format': {
                                'font_name': run.font.name,
                                'font_size': run.font.size,
                                'bold': run.font.bold,
                                'italic': run.font.italic,
                                'underline': run.font.underline,
                                'color': run.font.color.rgb if run.font.color.rgb else None,
                                'highlight': run.font.highlight_color
                            }
                        })
            current_pos = run_end
        
        for _ in range(len(paragraph.runs)):
            paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)
        
        for run_data in new_runs:
            new_run = paragraph.add_run(run_data['text'])
            fmt = run_data['format']
            
            if fmt['font_name']:
                new_run.font.name = fmt['font_name']
            if fmt['font_size']:
                new_run.font.size = fmt['font_size']
            if fmt['bold'] is not None:
                new_run.font.bold = fmt['bold']
            if fmt['italic'] is not None:
                new_run.font.italic = fmt['italic']
            if fmt['underline'] is not None:
                new_run.font.underline = fmt['underline']
            if fmt['color']:
                new_run.font.color.rgb = fmt['color']
            if fmt['highlight']:
                new_run.font.highlight_color = fmt['highlight']

def replace_placeholders_in_table(table, data_dict):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholders_in_paragraph(paragraph, data_dict)

def process_word_template(doc_bytes, data_dict):
    try:
        doc_io = io.BytesIO(doc_bytes)
        doc = Document(doc_io)
        for paragraph in doc.paragraphs:
            replace_placeholders_in_paragraph(paragraph, data_dict)
        for table in doc.tables:
            replace_placeholders_in_table(table, data_dict)
        return doc
    except Exception as e:
        st.error(f"Lỗi khi xử lý template Word: {str(e)}")
        return None

def convert_docx_to_pdf_libreoffice(docx_path, pdf_path):
    """
    Convert bằng LibreOffice với cấu hình tối ưu cho Streamlit Cloud
    """
    try:
        # Kiểm tra LibreOffice có sẵn không
        check_cmd = ['libreoffice', '--version']
        subprocess.run(check_cmd, capture_output=True, timeout=5)
        
        # Convert với options tối ưu
        cmd = [
            'libreoffice',
            '--headless',
            '--invisible',
            '--nocrashreport',
            '--nodefault',
            '--nofirststartwizard',
            '--nolockcheck',
            '--nologo',
            '--norestore',
            '--convert-to', 'pdf:writer_pdf_Export',
            '--outdir', os.path.dirname(pdf_path),
            docx_path
        ]
        
        env = os.environ.copy()
        env['HOME'] = tempfile.gettempdir()
        
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=120,
            env=env
        )
        
        # Xử lý tên file output
        expected_pdf = os.path.join(
            os.path.dirname(pdf_path),
            os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
        )
        
        if os.path.exists(expected_pdf) and expected_pdf != pdf_path:
            os.rename(expected_pdf, pdf_path)
        
        if not os.path.exists(pdf_path):
            raise Exception(f"PDF không được tạo. LibreOffice output: {result.stderr}")
        
        return True
        
    except FileNotFoundError:
        raise Exception(
            "⚠️ LibreOffice chưa được cài đặt!\n\n"
            "Cài đặt:\n"
            "• Local: sudo apt-get install libreoffice-writer\n"
            "• Streamlit Cloud: Thêm vào packages.txt:\n"
            "  libreoffice\n"
            "  libreoffice-writer"
        )
    except subprocess.TimeoutExpired:
        raise Exception("Timeout khi convert. File quá lớn hoặc phức tạp.")
    except Exception as e:
        raise Exception(f"Lỗi convert: {str(e)}")

def create_output_files(template_bytes, excel_data, selected_columns):
    output_files = []
    pdf_files = []
    
    # Tạo thư mục tạm
    tmpdir = tempfile.mkdtemp()
    
    try:
        # Tạo file Word
        for index, row in excel_data.iterrows():
            data_dict = {col: row[col] if pd.notna(row[col]) else "" for col in selected_columns}
            doc = process_word_template(template_bytes, data_dict)
            
            if doc is not None:
                # Tên file
                filename = f"output_{index + 1}.docx"
                for key in ['name', 'Name', 'ho_ten', 'ten', 'fullName', 'FullName', 'StudentName']:
                    if key in data_dict and data_dict[key]:
                        safe_name = re.sub(r'[^\w\s-]', '', str(data_dict[key]))
                        filename = f"{safe_name}.docx"
                        break
                
                docx_path = os.path.join(tmpdir, filename)
                doc.save(docx_path)
                
                with open(docx_path, "rb") as f:
                    output_files.append((filename, f.read()))
        
        # Convert sang PDF với progress bar
        if output_files:
            st.write("🔄 Đang chuyển đổi sang PDF...")
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            for idx, (filename, _) in enumerate(output_files):
                docx_path = os.path.join(tmpdir, filename)
                pdf_filename = filename.replace(".docx", ".pdf")
                pdf_path = os.path.join(tmpdir, pdf_filename)
                
                try:
                    status_text.text(f"Đang xử lý: {filename} ({idx+1}/{len(output_files)})")
                    convert_docx_to_pdf_libreoffice(docx_path, pdf_path)
                    
                    with open(pdf_path, "rb") as f:
                        pdf_files.append((pdf_filename, f.read()))
                        
                except Exception as e:
                    st.warning(f"⚠️ Không thể convert {filename}: {str(e)}")
                
                progress_bar.progress((idx + 1) / len(output_files))
            
            status_text.empty()
            progress_bar.empty()
    
    finally:
        # Cleanup temp files
        try:
            import shutil
            shutil.rmtree(tmpdir)
        except:
            pass
    
    return output_files, pdf_files

def create_zip_file(output_files):
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for filename, file_content in output_files:
            zip_file.writestr(filename, file_content)
    zip_buffer.seek(0)
    return zip_buffer.getvalue()

def merge_pdfs(pdf_contents):
    merger = PdfMerger()
    for _, pdf_data in pdf_contents:
        merger.append(io.BytesIO(pdf_data))
    output_buffer = io.BytesIO()
    merger.write(output_buffer)
    merger.close()
    output_buffer.seek(0)
    return output_buffer.getvalue()

# ---------------------- MAIN APP ----------------------
st.set_page_config(
    page_title="Tạo Word từ Excel", 
    page_icon="📄", 
    layout="wide",
    initial_sidebar_state="expanded"
)

st.title("📄 Tạo Word từ Excel & In hàng loạt")
st.markdown("---")

# Sidebar
with st.sidebar:
    st.header("📁 Upload Files")
    excel_file = st.file_uploader("Chọn file Excel (.xlsx, .xls)", type=['xlsx', 'xls'])
    word_file = st.file_uploader("Chọn file Word template (.docx)", type=['docx'])
    
    st.markdown("---")
    
    with st.expander("💡 Hướng dẫn sử dụng"):
        st.markdown("""
        **Bước 1:** Upload file Excel chứa dữ liệu
        
        **Bước 2:** Upload file Word template với placeholder {{tên_cột}}
        
        **Bước 3:** Chọn cột cần điền
        
        **Bước 4:** Nhấn "Tạo Files"
        """)
    
    with st.expander("⚙️ Cấu hình cho Streamlit Cloud"):
        st.code("""
# packages.txt
libreoffice
libreoffice-writer

# requirements.txt
streamlit
pandas
python-docx
openpyxl
PyPDF2
        """, language="text")
    
    with st.expander("✨ Tips giữ định dạng"):
        st.markdown("""
        - ✅ Dùng font: Arial, Times New Roman, Calibri
        - ✅ Tránh WordArt, effects phức tạp
        - ✅ Đặt margins: 2cm mỗi cạnh
        - ✅ Dùng styles có sẵn trong Word
        - ✅ Test template trước khi chạy hàng loạt
        """)

# Main content
if excel_file and word_file:
    excel_data = load_excel_file(excel_file)
    
    if excel_data is not None:
        template_bytes = word_file.getvalue()
        template_doc = Document(word_file)

        # Hiển thị dữ liệu
        st.subheader("📊 Dữ liệu Excel")
        st.dataframe(excel_data.head(10), use_container_width=True)
        st.caption(f"Tổng số dòng: {len(excel_data)}")

        # Chọn cột
        selected_columns = st.multiselect(
            "🎯 Chọn cột làm placeholder",
            options=excel_data.columns.tolist(),
            default=excel_data.columns.tolist(),
            help="Chọn các cột sẽ được điền vào template Word"
        )

        # Tìm placeholder
        placeholders = set()
        for paragraph in template_doc.paragraphs:
            placeholders.update(re.findall(r'\{\{([^}]+)\}\}', paragraph.text))
        for table in template_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        placeholders.update(re.findall(r'\{\{([^}]+)\}\}', paragraph.text))

        if placeholders:
            st.subheader("🔍 Placeholder trong template:")
            cols = st.columns(4)
            for idx, placeholder in enumerate(sorted(placeholders)):
                with cols[idx % 4]:
                    st.code(f"{{{{{placeholder}}}}}", language="text")
        else:
            st.warning("⚠️ Không tìm thấy placeholder nào. Định dạng: {{tên_cột}}")

        # Nút tạo file
        if selected_columns:
            col1, col2 = st.columns([1, 3])
            with col1:
                create_btn = st.button("🎯 Tạo Files", type="primary", use_container_width=True)
            
            if create_btn:
                with st.spinner("⏳ Đang xử lý..."):
                    output_files, pdf_files = create_output_files(
                        template_bytes, 
                        excel_data, 
                        selected_columns
                    )

                    if output_files:
                        st.success(f"✅ Hoàn thành! {len(output_files)} Word | {len(pdf_files)} PDF")
                        
                        # Download buttons
                        col1, col2, col3 = st.columns(3)
                        
                        with col1:
                            zip_content = create_zip_file(output_files)
                            st.download_button(
                                label="📦 Tải Word (.zip)",
                                data=zip_content,
                                file_name="word_documents.zip",
                                mime="application/zip",
                                use_container_width=True
                            )

                        with col2:
                            if pdf_files:
                                pdf_zip_content = create_zip_file(pdf_files)
                                st.download_button(
                                    label="📦 Tải PDF (.zip)",
                                    data=pdf_zip_content,
                                    file_name="pdf_documents.zip",
                                    mime="application/zip",
                                    use_container_width=True
                                )

                        with col3:
                            if pdf_files:
                                merged_pdf = merge_pdfs(pdf_files)
                                st.download_button(
                                    label="🖨️ PDF gộp (in)",
                                    data=merged_pdf,
                                    file_name="merged_output.pdf",
                                    mime="application/pdf",
                                    use_container_width=True
                                )

                        # Preview
                        if pdf_files and len(pdf_files) > 0:
                            st.subheader("👁️ Xem trước PDF đầu tiên")
                            first_pdf = pdf_files[0][1]
                            b64 = base64.b64encode(first_pdf).decode()
                            st.markdown(
                                f'<iframe src="data:application/pdf;base64,{b64}" width="100%" height="800px"></iframe>',
                                unsafe_allow_html=True
                            )
                    else:
                        st.error("❌ Không tạo được file nào. Kiểm tra lại template và dữ liệu.")
        else:
            st.warning("⚠️ Vui lòng chọn ít nhất một cột từ Excel")
else:
    # Welcome screen
    st.info("👆 **Bắt đầu:** Upload file Excel và Word template từ sidebar")
    
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("""
        ### 📝 File Excel cần có:
        - Header ở dòng đầu tiên
        - Dữ liệu từ dòng thứ 2 trở đi
        - Tên cột rõ ràng (vd: ho_ten, email, diem)
        """)
    
    with col2:
        st.markdown("""
        ### 📄 File Word template:
        - Sử dụng placeholder: `{{tên_cột}}`
        - Ví dụ: `Xin chào {{ho_ten}}`
        - Placeholder phải khớp với tên cột Excel
        """)
